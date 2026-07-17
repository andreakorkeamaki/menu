from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("output/documents/Trascrizione_integrale_Sartoria_Gastronomica.docx")
BURGUNDY = "7E2418"
GOLD = "A47B27"
PALE = "F5EEE9"
GRAY = "666666"


def P(name, description="", price="", allergens="", notes=""):
    return {"name": name, "description": description, "price": price, "allergens": allergens, "notes": notes}


classic = [
    P("SPRITZ APEROL", "Aperol, prosecco, soda", "6,50 €"),
    P("SPRITZ CAMPARI", "Campari, prosecco, soda", "6,50 €"),
    P("HUGO SPRITZ", "Sciroppo fiore di sambuco, prosecco, menta", "8 €"),
    P("AMERICANO", "Vermouth rosso, bitter Campari, soda", "8,00 €"),
    P("TOMMY’S MARGARITA", "Tequila, succo di lime fresco, sciroppo di agave", "10,00 €"),
    P("NEGRONI", "Gin, Vermouth rosso, bitter Campari", "9,00 €"),
    P("NEGRONI SBAGLIATO", "Vermouth rosso, bitter Campari, prosecco", "9,00 €"),
    P("CAMPARI ORANGE", "Campari Bitter, succo d’arancia", "8,00 €"),
    P("MI TO", "Campari Bitter, Vermouth", "8,00 €", notes="Nel menu Bakery Café la stessa voce è scritta “MITO”."),
    P("CUBA LIBRE", "Rum scuro, lime, coca cola", "10,00 €"),
    P("PALOMA", "Tequila, soda al pompelmo, lime", "10,00 €"),
    P("GIN TONIC", "Luis Gin 8, tonica Schweppes", "9,00 €"),
    P("GIN TONIC PREMIUM", "Tonica Indian Fever Tree e Gin da scegliere tra le nostre 60 etichette", "Da 10,00 a 15,00 €"),
    P("MOSCOW MULE", "Vodka, lime, ginger beer", "9,00 €"),
    P("MOJITO", "Rum, lime, zucchero di canna, menta", "10,00 €"),
    P("BOLOGNA Menegiks Gin & Tonic", "TONICA: Indian Tonic Fever Tree\nBOTANICHE: ginepro, acqua di mare, lime, menta piperita, cappero", "15,00 €"),
    P("BLOODY MARY", "Vodka, succo di pomodoro, lime, spezie", "10,00 €"),
    P("MANHATTAN", "Rye Whisky, Vermouth rosso, angostura", "10,00 €"),
    P("BOULEVARDIER", "Bitter Campari, Vermouth rosso, Bourbon Whisky", "10,00 €"),
    P("WHISKY SOUR", "Bourbon Whisky, succo fresco di limone, zucchero, albumina", "10,00 €"),
    P("DAIQUIRI", "Rhum, succo fresco di lime, zucchero", "10,00 €"),
    P("MARGARITA", "Tequila, Triple Sec, succo fresco di limone", "10,00 €"),
    P("MIMOSA", "Prosecco, succo d’arancia fresca", "10,00 €"),
    P("KIR", "Creme de cassis, vino bianco fermo", "10,00 €"),
    P("KIR ROYAL", "Creme de cassis, prosecco", "10,00 €"),
    P("DRY MARTINI", "Gin e Vermouth Dry", "10,00 €"),
    P("GIN FIZZ", "Gin, succo fresco di limone, sciroppo di zucchero, soda", "10,00 €"),
    P("ANALCOLICO FRUTTA", "Fantasia del barman", "8,00 €"),
    P("OLD CUBAN", "Rum scuro, lime, zucchero, angostura, prosecco, menta", "10,00 €", notes="Disponibile dalle 18:00."),
    P("OLD FASHIONED", "Bourbon whiskey, zucchero, angostura", "10,00 €", notes="Disponibile dalle 18:00."),
    P("SAZERAC", "Cognac, assenzio, zucchero, peychaud, angostura", "10,00 €", notes="Disponibile dalle 18:00."),
    P("VIEUX CARRE", "Rye whiskey, cognac, vermouth, benedectine, peychaud bitter", "12,00 €", notes="Disponibile dalle 18:00."),
    P("WHITE LADY", "Gin, triple sec, limone", "10,00 €", notes="Disponibile dalle 18:00."),
    P("WHITE RUSSIAN", "Vodka, kalhua, schiuma di latte", "10,00 €", notes="Disponibile dalle 18:00."),
    P("PENCILLIN", "Whiskey, honey mix, limone", "12,00 €", notes="Disponibile dalle 18:00. Grafia della fonte."),
    P("CARDINALE", "Campari bitter, vermouth dry, gin", "9,00 €", notes="Disponibile dalle 18:00."),
    P("HANKY PANKY", "Gin, vermouth, fernet", "10,00 €", notes="Disponibile dalle 18:00."),
    P("DARK’N’STORMY", "Rum scuro, lime, ginger beer", "10,00 €", notes="Disponibile dalle 18:00."),
    P("MINT JULEP", "Whiskey, zucchero, menta", "10,00 €", notes="Disponibile dalle 18:00."),
    P("TOM COLLINS", "Gin old tom, limone, zucchero, soda", "10,00 €", notes="Disponibile dalle 18:00."),
    P("GOD FATHER", "Scotch whiskey, amaretto", "10,00 €", notes="Disponibile dalle 18:00."),
    P("GOD MOTHER", "Vodka, amaretto", "10,00 €", notes="Disponibile dalle 18:00."),
    P("TRINIDAD SOUR", "Rye whiskey, orzata, limone, angostura", "10,00 €", notes="Disponibile dalle 18:00."),
    P("JUNGLE BIRD", "Rum scuro, bitter, ananas, zucchero, lime", "10,00 €", notes="Disponibile dalle 18:00."),
]

signature = [
    P("CINNAMON CAL", "Bitter Gamondi infuso all’ananas\nRum Copalli al cacao\nSciroppo di zucchero e cannella handmade\nSucco di limone\nTecnica: Build\nIntensità: YYY", "12,00 €"),
    P("PEPE AR CU...IWI", "Rum scuro\nAmaro Ilex\nSciroppo di kiwi homemade\nScorza di limone\nPepe verde\nSeltz\nTecnica: Build\nIntensità: YY", "12,00 €", notes="Titolo riportato esattamente come appare nella fonte; la parte sostituita dai puntini non è leggibile."),
    P("SEDANO DI VASINICOLA", "Mezcal\nBasilico fresco\nSciroppo d’agave\nLime\nCrosta di sale, pepe e sedano essicato\nTecnica: Shake and strain\nIntensità: YY", "12,00 €"),
    P("ARANCIA MOSCATA", "Vermouth\nVaniglia\nNoce moscata\nArancia fresca\nZucchero\nAngostura all’arancia\nTecnica: Shake and strain\nIntensità: YY", "12,00 €"),
    P("GOLDEN FIZZ", "Gin\nLimone\nPepe\nZucchero moscovado\nCurcuma\nSeltz\nTecnica: Shake and strain + Build\nIntensità: YY", "12,00 €"),
    P("RED PEPPER LADY", "Vodka\nSciroppo di ibisco\nLimone\nPepe rosa\nTriple sec\nTecnica: Shake and strain\nIntensità: YY", "12,00 €"),
    P("ANICE ROYALE", "Shrub di melograno\nAnice stellato\nProsecco\nAngostura fiorita", "12,00 €"),
    P("CANELLE REPONDE", "Vermouth analcolico\nSucco di pesca e mela\nCannella\nZucchero\nTecnica: Shake and strain\nIntensità: YY\nAlcool Free", "10,00 €"),
    P("SOFY RAMONES", "Sciroppo al Karcadè\nPepe rosa\nDistillato analcolico\nLimone\nTecnica: Build\nIntensità: YYY\nAlcool Free", "10,00 €"),
]

gin_tonic = [
    P("TOSCANA Sabatini & Tonic", "TONICA: Indian Tonic Fever Tree\nBOTANICHE: ginepro, iris fiorentino, coriandolo, foglie di olivo, salvia, timo, lavanda, finocchietto selvatico, lemon verbena", "12,00 €"),
    P("TOSCANA Peter in Florence & Tonic", "TONICA: Indian Tonic Fever Tree\nBOTANICHE: ginepro, iris fiorentino, scorza di bergamotto fresca, scorza di limone essiccata, bacche di rosa, fiori di lavanda e di rosmarino freschi, radice di angelica, coriandolo, mandorle amare. Tutti prodotti provenienti da micro produttori locali specializzati e bio", "15,00 €"),
    P("TOSCANA Peter in the Navy & Tonic", "TONICA: Indian Tonic Fever Tree\nBOTANICHE: ginepro, radice di iris, petali di iris, scorza di limone essiccata, scorza di bergamotto, semi di cardamomo, radice di angelica, coriandolo, corteccia di cassia", "15,00 €"),
    P("TOSCANA 1592 & Tonic", "TONICA: Indian Tonic Fever Tree\nBOTANICHE: Mono botanica ginepro", "12,00 €"),
    P("BASSA MODENA Tabar & Tonic", "TONICA: Indian Tonic Fever Tree\nBOTANICHE: Ginepro, Coriandolo, Angelica, Arancia Amara, Camomilla, Semi di Anice, Cardamomo, Arancia Dolce, Rosmarino", "12,00 €"),
    P("BOLOGNA Seawake & Tonic", "TONICA: Indian Tonic Fever Tree\nBOTANICHE: Ginepro, Cardamomo, coriandolo, finocchio, scorze di arancia", "12,00 €"),
    P("BOLOGNA Bad Gin & Tonic", "TONICA: Indian Tonic Fever Tree\nBOTANICHE: Ginepro, coriandolo, limone, pepe nero", "12,00 €"),
    P("BOLOGNA Nettuno Blu & Tonic", "TONICA: Indian Tonic Fever Tree\nBOTANICHE: Ginepro, camomilla, alga spirulina, lavanda, lime, pepe di Sichuan, zenzero, rosa, coriandolo", "12,00 €"),
    P("BOLOGNA Nettuno Dry & Tonic", "TONICA: Indian Tonic Fever Tree\nBOTANICHE: Ginepro toscano, angelica, radice di giaggiolo, cardamomo nero, bergamotto, limone, pepe, cannella regina", "10,00 €"),
]

cocktail_beers = [
    P("MORETTI BIONDA 4.6%", "0.30 / 0.40", "4,00 € / 5,00 €"),
    P("MORETTI IPA 5.0%", "0.30 / 0.40", "5,00 € / 6,00 €"),
    P("MORETTI BIANCA 5.0%", "0.30 / 0.40", "5,00 € / 6,00 €"),
    P("MORETTI 0.0 ALCOOL", "Analcolica", "5,00 €"),
    P("LA GRIGNA PILS", "Birrificio Lariano 33 cl 6.6%", "6,00 €"),
    P("LATEX PIÙ BLANCHE", "Birrificio Reporto 33 cl 4.8%", "6,00 €"),
    P("JAKE PACIFIC PALE ALE", "Gluten Free – Birrificio Bellazzi 33 cl 5.5%", "6,00 €"),
    P("MY BO RED", "Birrificio Bellazzi 33 cl 5.0%", "6,00 €"),
    P("AMERICAN HYPE IPA", "Birrificio Lariano 33 cl 6.8%", "7,00 €"),
    P("WEISS", "Birrificio Manifattura 33 cl 5.0%", "6,00 €"),
]

stuzz = [
    P("TAGLIERE SALUMI ASSORTITI", "Per 2 persone", "14,00 €"), P("TAGLIERE SALUMI ASSORTITI", "Per 4 persone", "28,00 €"),
    P("POLPETTE DI MANZO", "Alle tre salse", "10,00 €"), P("POLPETTE DI PESCE SPADA FRITTE", "con maionese al limone", "16,00 €"),
    P("POLPETTE MELANZANE E MENTA", "Con salsa tzatziki", "10,00 €"), P("POLPETTE DI LENTICCHIE", "Su specchio di pomodoro", "10,00 €"),
    P("COPPO SICILIANO", "Panelle – Arancinette – Crocchette", "12,00 €"), P("COPPO CALAMARI E GAMBERI", "", "16,00 €"),
    P("CESTINO DI SPICCHI DI PIADA", "", "5,00 €"), P("TIGELLA ARTIGIANALE", "", "1,50 €"),
]

lunch_notes = [
    "Menù valido dal lunedì al venerdì",
    "Menù da una portata con primo piatto o insalatona € 13",
    "Menù da una portata con secondo piatto o panino gigante € 16",
    "Menù doppia portata (non condivisibile) € 30",
    "Sempre compresi 1/2 litro di acqua, un calice di vino e caffè",
]
lunch_primi = [
    P("PRIMO DEL GIORNO", "Chiedere allo staff", "13,00 €"),
    P("LASAGNA SFOGLIA VERDE ALLA BOLOGNESE", price="13,00 €", allergens="Cereali – Glutine (1), Uova (3), Latte (7), Sedano (9), Anidride solforosa e solfiti (12)"),
    P("PARMIGIANA DI MELANZANE", price="13,00 €", allergens="Cereali – Glutine (1), Uova (3), Latte (7)"),
    P("ZUPPA DI LEGUMI E CEREALI", "Disponibile in autunno/inverno", "13,00 €", "Glutine (1)"),
    P("PITA GRECA", "Con pollo sfilacciato, cipolle, olive kalamata, feta, iceberg, pomodoro, salsa tzatziki e patate arrosto", "15,00 €", "Glutine (1), Latticini (7)"),
]
lunch_secondi = [
    P("PETTO DI POLLO ALLA GRIGLIA", "Sale grosso e rosmarino con patate al rosmarino\n(gluten free)", "16,00 €"),
    P("BISTECCA DI MANZO ALLA GRIGLIA", "Con patate al rosmarino\n(gluten free)", "16,00 €"),
    P("PESCE DEL GIORNO", "Con verdure", "16,00 €"),
    P("BOMBETTE DI MAIALE", "Con ricotta e pistacchio", "16,00 €", "Cereali – Glutine (1), Latte (7), Frutta a guscio (8)"),
    P("SPIEDINI DI POLLO SPEZIATO", "Con patate al rosmarino\n(gluten free)", "16,00 €"),
    P("SPIEDINO TURCO", "Con carne di manzo macinata, con feta, cipolle, olive Kalamata e pita", "16,00 €", "Latte (7)\nGlutine (1), Latticini (7)", "La fonte mostra due righe separate introdotte da “Allergeni”."),
    P("SPECIAL DEL GIORNO", "Chiedere allo staff", "16,00 €"),
    P("MEZZO POLLO AL FORNO", "Con patate e rosmarino", "16,00 €"),
]
lunch_salads = [
    P("INSALATA DI POLLO", "Con petto di pollo alla griglia\n(gluten free)", "13,00 €"),
    P("CAESAR SALAD", "Con iceberg, scaglie di parmigiano, dressing, petto di pollo e crostini", "13,00 €", "Cereali – Glutine (1), Uova (3), Latte (7), Anidride solforosa (12)"),
    P("INSALATA DI MANZO", "Con straccetti di manzo alla griglia\n(gluten free)", "13,00 €"),
    P("PANZANELLA CROCCANTE", "Mozzarella, olive kalamata, pomodorini, crostini di pane e basilico", "13,00 €", "Glutine (1), Latticini (7)"),
    P("PANZANELLA CROCCANTE VEG", "Tofu al basilico, olive Kalamata, pomodorini, crostini di pane, basilico", "13,00 €", "Glutine (1), Latticini (7)"),
    P("SALMONE MARINATO", "Iceberg, finocchio, pomodoro e crostini", "13,00 €", "Cereali – Glutine (1), Pesce (4)"),
    P("GRECA", "Con feta, origano, cetrioli, cipolla di Tropea e olive Kalamata (vegetariana)", "13,00 €", "Latte (7)"),
    P("PEPERONI RIPIENI", "Con tonno, insalata mista, olive e crostini di pane", "13,00 €", "Glutine (1), Anidride solforosa (12)"),
    P("NOCI E GORGONZOLA", "Con insalata mista", "13,00 €", "Frutta a guscio, Latticini (7)"),
    P("SPECK, CARCIOFI E OLIVE", "Con carciofi sott’olio e insalata mista", "13,00 €"),
    P("POLLO SFILACCIATO", "Con insalata mista e crostini di pane", "13,00 €", "Glutine (1), Latticini (7)"),
    P("TONNO E AVOCADO", "Con aceto balsamico e insalata mista", "13,00 €", "Pesce, Anidride solforosa (12)"),
]
lunch_extras = [
    P("MOZZARELLA DI BUFALA", price="2,50 €"), P("FETA", price="2,50 €"), P("AVOCADO", price="3,50 €"),
    P("SALMONE MARINATO", price="3,50 €"), P("POLLO STRACCETTI", price="3,50 €"), P("PANINO SENZA GLUTINE", price="2,50 €"),
]
lunch_burgers = [
    P("CLASSICO", "Con maionese, ketchup, insalata e pomodoro", "16,00 €", "Cereali – Glutine (1), Uova (3), Semi di sesamo (11)"),
    P("SARTORIA BURGER", "Con bacon, cheddar, maionese, insalata e pomodoro", "16,00 €", "Cereali – Glutine (1), Uova (3), Latte (7), Semi di sesamo (11)"),
    P("SMASHED BURGER MANZO", "Burger schiacciato sulla piastra, cheddar, bacon, salsa hamburger, e patatine fritte", "16,00 €", "Glutine (1), Uova (3), Latte (7), Semi di sesamo (11)"),
    P("SMASHED BURGER POLLO", "Burger schiacciato sulla piastra, cheddar, bacon, salsa hamburger, e patatine fritte", "16,00 €", "Glutine (1), Uova (3), Latte (7), Semi di sesamo (11)"),
    P("PANINO CON BISTECCA", "Di manzo, maionese, insalata e pomodoro", "16,00 €", "Cereali – Glutine (1), Uova (3), Latte (7), Semi di sesamo (11)"),
    P("VEGETARIANO", "Maionese, insalata e pomodoro", "16,00 €", "Cereali – Glutine (1), Uova (3), Semi di sesamo (11)"),
]
lunch_desserts = [
    P("NY CHEESECAKE", "Con frutti di bosco\n(prodotto da noi)", "5,00 €", "Cereali – Glutine (1), Latte (7), Uova (3)"),
    P("NY CHEESECAKE", "Con cioccolato\n(prodotto da noi)", "5,00 €", "Cereali – Glutine (1), Latte (7), Uova (3)"),
    P("TENERINA", "Con panna montata (prodotto da noi)", "5,00 €", "Uova (3), Latte (7)"),
    P("TORTA DI MELE", "Con panna montata (prodotto da noi)", "5,00 €", "Glutine (1), Uova (3), Latte (7)"),
    P("GELATO ALLA CREMA", "", "4,00 €", "Uova (3), Latte (7)"),
    P("AFFOGATO", "Al caffè o liquori", "5,00 €", "Uova (3), Latte (7)"),
]

bakery_pastry = [
    P("PASTICCERIA MIGNON", "Bignè, Cestino alla frutta, Cannolo, Tiramisù, Cheesecake, etc.", "Da 1,50 a 5,00 €"),
    P("CROISSANT", "Vuoto o Ripieno (crema, marmellata, cioccolato, ricotta, pistacchio, etc.)", "Da 1,80 a 3,50 €"),
    P("TENERINA AL CIOCCOLATO", "Torta morbida al cioccolato", "6,00 €"), P("APPLE PIE", "Torta di mele", "6,00 €"),
    P("NY CHEESECACKE", "Ai frutti rossi", "6,00 €", notes="Grafia della fonte."), P("CHEESECACKE", "Al cioccolato", "6,00 €", notes="Grafia della fonte."),
]
bakery_cafe = [
    P("CAFFÈ ESPRESSO", price="2,00 €"), P("CAPPUCCINO", "+ latte vegetale (soia, avena, HD) +0,50 €", "2,50 €"),
    P("CAFFÈ AMERICANO", price="2,50 €"), P("ORZO", "Piccolo / Grande", "2,30 € / 3,30 €"),
    P("GINSENG", "Piccolo / Grande", "2,50 € / 3,50 €"), P("TÈ CALDO", price="4,00 €"),
]
bakery_special = [
    P("CREMA CAFFÈ", price="3,50 €"), P("NERO DI BRONTE", "Salsa al pistacchio, caffè, crema di latte montata a caldo", "4,50 €"),
    P("CAFFÈ SHAKERATO", price="3,50 €"), P("CAFFÈ ESTIVO", price="3,70 €"), P("CAFFÈ SALENTINO", price="3,00 €"), P("CAFFÈ IN GHIACCIO", price="2,50 €"),
]
bakery_summer = [
    P("CREMA DI CAFFÈ", price="3,50 €"), P("FROZEN YOGURT AL CIOCCOLATO", price="6,00 €"),
    P("FROZEN YOGURT CON AMARENE E MERINGA", price="6,00 €"), P("FROZEN YOGURT CON CARAMELLO", price="6,00 €"),
]
bakery_soft = [
    P("SPREMUTA D’ARANCIA", price="5,50 €"), P("SUCCHI DI FRUTTA", price="4,00 €"), P("BIBITE IN BOTTIGLIA 33CL", price="4,00 €"),
    P("PICCOLI APERITIVI", price="4,00 €"), P("BOLLICINE", price="6,00 € / 8,00 €"), P("CALICE DI VINO", price="6,00 € / 8,00 €"), P("SANGRIA", price="5,00 €"),
]
bakery_classic = [dict(x) for x in classic[:15]]
bakery_classic[8]["name"] = "MITO"
bakery_beers = [
    P("MORETTI AL MELOGRANO 5%", price="5,00 €"), P("MORETTI BIONDA 4.6%", "0.30 / 0.40", "4,00 € / 5,00 €"),
    P("MORETTI IPA 5.0%", "0.30 / 0.40", "5,00 € / 6,00 €"), P("MORETTI IPA 5.0%", "0.30 / 0.40", "5,00 € / 6,00 €", notes="Duplicato identico presente nella fonte."),
    P("MORETTI 0.0 ALCOOL", "Analcolica", "5,00 €"), P("LA GRIGNA PILS", "Birrificio Lariano 33 cl 6.6%", "6,00 €"),
    P("LATEX PIÙ BLANCHE", "Birrificio Reporto 33 cl 4.8%", "6,00 €"), P("JAKE PACIFIC PALE ALE", "Birrificio Bellazzi 33 cl 5.5% – Gluten Free", "6,00 €"),
    P("MY BO RED", "Birrificio Bellazzi 33 cl 5.0%", "6,00 €"),
]

dinner_antipasti = [
    P("CRUDO DI PARMA", "Stagionato 24 mesi – Selezione Villani", "16,00 €"),
    P("TAGLIERE DI SALUMI ASSORTITI PER 2 PERSONE", "Con mortadella IGP del consorzio (selezione Villani) e squacquerone", "14,00 €", "Latte (7)"),
    P("TAGLIERE DI SALUMI ASSORTITI PER 4 PERSONE", "Con mortadella IGP del consorzio (selezione Villani) e squacquerone", "28,00 €", "Latte (7)"),
    P("CESTINO DI SPICCHI DI PIADINE ARTIGIANALI", price="5,00 €", allergens="Cereali – Glutine (1)"), P("TIGELLA ARTIGIANALE", price="1,50 €", allergens="Glutine (1)"),
    P("BATTUTA AL COLTELLO DI FASSONA", "Con uovo fritto e maionese all’acciuga", "25,00 €", "Uova (3), Pesce (4)"),
    P("CARPACCIO DI MANZO", "Con crema di parmigiano e pomodorini confit", "18,00 €", "Latte (7)"),
    P("SFORMATINO DI VERDURE", "Con chantilly di Robiola d’Alba DOP", "14,00 €", "Uova (3), Latte (7)"),
    P("BURRATA CON CICORIA", "Ripassata al peperoncino", "16,00 €", "Latte (7)"), P("BURRATA CON CRUDO DI PARMA", price="22,00 €", allergens="Latte (7)"),
    P("SELEZIONE DI SALUMI E FORMAGGI", "Accompagnati da confetture", "16,00 €", "Latte (7)"),
    P("SFORMATINO DI PATATE E MORTADELLA", "Con cuore morbido di caciotta", "15,00 €", "Latte (7)"),
    P("POLPETTE DI ZUCCA E PATATE", "Su spinaci freschi con crema di grana", "14,00 €", "Cereali – Glutine (1), Uova (3), Latte (7)"),
    P("POLPETTE DI MELANZANE E MENTA FRITTE", "Con salsa tzatziki", "14,00 €", "Cereali – Glutine (1), Uova (3), Latte (7)"),
    P("POLPETTE DI LENTICCHIE", "Affogate al pomodoro e basilico", "14,00 €", "Cereali – Glutine (1), Uova (3), Latte (7), Sedano"),
    P("POLPETTE DI PESCE SPADA FRITTE", "Su letto di rucola e maionese al limone", "16,00 €", "Cereali – Glutine (1), Uova (3), Pesce"),
    P("COCKTAIL DI GAMBERI", "Anni 80", "17,00 €", "Cereali – Pesce"),
]
dinner_primi = [
    P("TORTELLINI FATTI A MANO", "Con ragù bianco su crema di Parmigiano Reggiano", "17,00 €", "Cereali – Glutine (1), Uova (3), Latte (7)"),
    P("TORTELLINI FATTI A MANO IN BRODO", price="15,00 €", allergens="Cereali – Glutine (1), Uova (3), Latte (7)"),
    P("TORTELLINO ORO DI BOLOGNA", "Con crema di panna, parmigiano e zafferano", "16,00 €", "Cereali – Glutine (1), Uova (3), Latte (7)"),
    P("QUADRATONI DI MELANZANE, SCAMORZA E RICOTTA", "Con vellutata di pomodoro e basilico", "16,00 €", "Cereali – Glutine (1), Uova (3), Latte (7)"),
    P("RAVIOLI RIPIENI DI SALSICCIA, PEPERONI E RICOTTA", "Con ragù di salsiccia fatta da noi", "16,00 €", "Cereali – Glutine (1), Uova (3), Latte (7)"),
    P("RAVIOLI DI RICOTTA, ’NDUJA E CACIOCAVALLO", "Su crema di parmigiano e pomodorini", "16,00 €", "Cereali – Glutine (1), Uova (3), Latte (7)"),
]
dinner_meat = [
    P("FIORENTINA", "T-Bone", "Da 68,00 €/kg a seconda del periodo di frollatura e provenienza"),
    P("COSTATA", "Rib Eye", "Da 58,00 €/kg a seconda del periodo di frollatura e provenienza"),
]
dinner_secondi = [
    P("FILETTO DI FASSONA ALLA GRIGLIA", price="26,00 €"),
    P("TAGLIATA NEW YORK TOP QUALITY", "Con fonduta di gorgonzola e noci", "25,00 €", "Frutta a guscio (8), Latte (7)"),
    P("TAGLIATA NEW YORK TOP QUALITY", "Con pesto di peperoncini", "24,00 €"),
    P("TAGLIATA NEW YORK TOP QUALITY", "Con scaglie di pecorino e granella di pistacchio", "26,00 €", "Latte (7), Frutta a guscio (8)"),
    P("TAGLIATA NEW YORK TOP QUALITY", "Con lardo di Patanegra", "24,00 €"),
    P("COSTOLA DI SUINO CARAMELLATA", "Con salsa BBQ e miele alla griglia", "18,00 €", "Anidride solforosa e solfiti (12)"),
]
dinner_sides = [
    P("PATATE GRATINATE", price="8,00 €", allergens="Glutine (1), Latte (7)"), P("PATATE ARROSTO AL SALE GROSSO E ROSMARINO", price="6,00 €"),
    P("PATATE FRITTE*", price="6,00 €"), P("CHIPS DI PATATE SALE E ROSMARINO", price="6,00 €"),
    P("SPINACI FRESCHI SALTATI CON OLIO EVO E PEPE", price="7,00 €"), P("VERDURE DI STAGIONE ALLA GRIGLIA", price="8,00 €"),
    P("FUNGHI PORCINI FRITTI*", price="15,00 €"), P("CICORIA SALTATA IN PADELLA", "Con aglio, olio e peperoncino", "7,00 €"),
]
dinner_salads = [
    P("SPECIAL SARTORIA POLLO", "Insalata mista con petto di pollo alla griglia e scaglie di Parmigiano Reggiano", "15,00 €", "Latte (7)"),
    P("SPECIAL SARTORIA MANZO", "Insalata mista con bistecca di manzo alla griglia e scaglie di Parmigiano Reggiano", "16,00 €", "Latte (7)"),
    P("INSALATA DI SALMONE", "Insalata mista, salmone marinato, finocchio, semi di girasole, Philadelphia, semi di zucca, salsa tzatziki, crostini", "15,00 €", "Cereali – Glutine (1), Pesce (4), Latte (7)"),
    P("INSALATA AVOCADO E GAMBERI", "Insalata mista, avocado, gamberi, mousse di Philadelphia e semi di girasole", "17,00 €", "Crostacei (2)"),
    P("INSALATA GRECA", "Insalata mista con cetrioli, feta, cipolla di Tropea, olive kalamata e origano", "15,00 €", "Latte (7)"),
    P("PANZANELLA VEGANA", "Crostini di pane aromatizzato al basilico, pomodorini, olive kalamata e tofu al basilico", "14,00 €", "Cereali – Glutine (1)"),
]
dinner_burgers = [
    P("SMASHED BURGER DI POLLO", "Petto di pollo, cheddar, bacon, salsa hamburger con patate fritte\nSCHIACCIATO SULLA PIASTRA!", "15,00 €", "Semi di sesamo (11), Cereali – Glutine (1), Latte (7), Uova (3)"),
    P("SMASHED BURGER DI MANZO", "Burger di manzo (200g), cheddar, bacon e salsa burger con patate fritte\nSCHIACCIATO SULLA PIASTRA!", "15,00 €", "Semi di sesamo (11), Cereali – Glutine (1), Latte (7), Uova (3)"),
    P("ALTA SARTORIA", "Doppio burger di manzo (400g), pancetta, cheddar e salsa burger", "22,00 €", "Semi di sesamo (11), Cereali – Glutine (1), Latte (7), Uova (3)"),
    P("BOLOGNESE", "Burger di manzo (200 gr), crema di parmigiano, mortadella (selezione Villani) a volontà, topping al balsamico", "20,00 €", "Semi di sesamo (11), Cereali – Glutine (1), Latte (7), Uova (3)"),
    P("MODENESE", "Burger di manzo (200g), pecorino, maionese al balsamico e glassa di aceto balsamico con patate fritte", "15,00 €", "Semi di sesamo (11), Anidride solforosa e solfiti (12), Cereali – Glutine (1), Latte (7), Uova (3)"),
    P("PICCANTE", "Burger di manzo (200g), cheddar, pancetta croccante e salsa piccante (maionese e pesto di peperoncino calabrese)", "16,00 €", "Cereali – Glutine (1), Latte (7), Uova (3), Semi di sesamo (11)"),
]
dinner_desserts = [
    P("NEW YORK CHEESECAKE", "Ai frutti di bosco o al cioccolato", "8,00 €", "Cereali – Glutine (1), Uova (3), Latte (7)"),
    P("FONDENTE LEGGERO", "Con crema chantilly e panna montata", "7,00 €", "Latte (7)"),
    P("TIRAMISÙ AL BICCHIERE", "Con mascarpone e caffè", "8,00 €", "Cereali – Glutine (1), Latte (7), Uova (3)"),
    P("MILLEFOGLIE SCOMPOSTA", "Con chantilly e frutta", "8,00 €", "Cereali – Glutine (1), Latte (7), Uova (3)"),
    P("TORTINO ALLE NOCCIOLE E CIOCCOLATO", "Con crema alla gianduia", "7,00 €", "Cereali – Glutine (1), Uova (3), Frutta a guscio (8)"),
    P("TORTA DI MELE E CANNELLA", "Servita con gelato alla crema e panna montata", "8,00 €", "Cereali – Glutine (1), Uova (3), Latte (7)"),
]
dinner_drinks = [P("ACQUA PANNA", "Minerale 0,75 lt", "3,00 €"), P("ACQUA SAN PELLEGRINO", "Minerale 0,75 lt", "3,00 €"), P("BIBITE IN BOTTIGLIA", "33 cl", "4,00 €")]
dinner_beers = [dict(x) for x in cocktail_beers]
dinner_beers[0]["name"] = "MORETTI BIONDA 5.0%"
dinner_wines = [
    P("BOLLICINE AL CALICE", "A partire da", "6,00 €"), P("BIANCHI AL CALICE", "A partire da", "6,00 €"), P("ROSSI AL CALICE", "A partire da", "6,00 €"),
    P("QUINTESSENZA PIGNOLETTO FRIZZANTE MEDICI ERMETE", "Bollicine", "22,00 €"),
    P("PROSECCO SUPERIORE DI VALDOBBIADENE DOCG BRUT BELLENDA", "Bollicine", "25,00 €"),
    P("FRANCIACORTA BRUT DOCG FERGHETTINA", "Bollicine", "35,00 €"),
    P("FRANCIACORTA ROSÉ DOCG FERGHETTINA", "Bollicine", "40,00 €"),
    P("PIGNOLETTO PERDITEMPO TERRE ROSSE", "Bollicine", "22,00 €"),
    P("LUGANA CA DEI FRATI", "Vino Bianco", "40,00 €"),
]


sections = [
    ("COCKTAIL BAR", [
        ("I classici", "Cocktail Bar – Sartoria Gastronomica.pdf", classic, ["La fonte divide gli ultimi 16 cocktail con la nota: DISPONIBILI DALLE 18:00."]),
        ("Signature", "2Cocktail Bar – Sartoria Gastronomica.pdf", signature, []),
        ("Gin & Tonic", "3Cocktail Bar – Sartoria Gastronomica.pdf", gin_tonic, []),
        ("Birre", "4Cocktail Bar – Sartoria Gastronomica.pdf", cocktail_beers, []),
        ("Stuzzicherie", "5Cocktail Bar – Sartoria Gastronomica.pdf", stuzz, []),
    ]),
    ("MENÙ PRANZO", [
        ("Note generali del menù", "Tutte le 5 pagine del Menù Pranzo", [], lunch_notes),
        ("Primi", "Menù Pranzo – Sartoria Gastronomica.pdf", lunch_primi, []),
        ("Secondi", "2Menù Pranzo – Sartoria Gastronomica.pdf", lunch_secondi, []),
        ("Insalatone", "3Menù Pranzo – Sartoria Gastronomica.pdf", lunch_salads, ["Tutte le insalatone hanno come base: iceberg, lattuga, radicchio, carote, pomodori"]),
        ("Aggiunte extra insalatone", "3Menù Pranzo – Sartoria Gastronomica.pdf", lunch_extras, []),
        ("Panini giganti", "4Menù Pranzo – Sartoria Gastronomica.pdf", lunch_burgers, ["Burger di manzo (200 gr) nel pan brioche ricoperto di semi di sesamo, prodotto da noi, servito con patate al rosmarino"]),
        ("Dessert", "5Menù Pranzo – Sartoria Gastronomica.pdf", lunch_desserts, []),
    ]),
    ("SARTORIA BAKERY CAFÈ", [
        ("Pasticceria", "Sartoria Bakery Cafè – Sartoria Gastronomica.pdf", bakery_pastry, []),
        ("Caffetteria", "Sartoria Bakery Cafè – Sartoria Gastronomica_1.pdf", bakery_cafe, []),
        ("Caffetteria speciale", "2Sartoria Bakery Cafè – Sartoria Gastronomica.pdf", bakery_special, []),
        ("Bevande estive", "3.pdf", bakery_summer, []),
        ("Soft drink", "S5artoria Bakery Cafè – Sartoria Gastronomica.pdf", bakery_soft, []),
        ("I classici", "6Sartoria Bakery Cafè – Sartoria Gastronomica.pdf", bakery_classic, ["DISPONIBILI DALLE 18:00 (nota mostrata in fondo alla pagina, dopo le voci visibili)."]),
        ("Birre", "7Sartoria Bakery Cafè – Sartoria Gastronomica.pdf", bakery_beers, ["Testi promozionali presenti nella pagina: L’APERITIVO CHE NON HAI MAI BEVUTO; Provala così; In calice; 3 cubetti di ghiaccio; foglia di menta; MESSINA NOTE DI MELOGRANO."]),
    ]),
    ("SARTORIA GASTRONOMICA — CENA", [
        ("Antipasti", "Sartoria Gastronomica – Sartoria Gastronomica.pdf", dinner_antipasti, []),
        ("Primi", "2Sartoria Gastronomica – Sartoria Gastronomica.pdf", dinner_primi, []),
        ("Tagliate su misura", "3Sartoria Gastronomica – Sartoria Gastronomica.pdf", dinner_meat, [
            "Selezione di carni della Sartoria a peso",
            "Carni a peso: le costate con prezzo ad ettogrammo vengono pesate al momento del taglio.",
            "Bisogna considerare che per la sua fisionomia un taglio costata non può essere inferiore ai 600/700 grammi mentre il taglio Fiorentina si aggira intorno al chilogrammo come minimo.",
            "La carne destinata alla frollatura (minimo 60 giorni) viene conservata in una cella di maturazione esposta nella sala ristorante.",
            "Chiedere al personale di sala le diverse razze e provenienze.",
        ]),
        ("Secondi", "4Sartoria Gastronomica – Sartoria Gastronomica.pdf", dinner_secondi, ["Tutti gli involtini e le bombette sono panati con pangrattato, pecorino e parmigiano accompagnati da insalatina.", "Nota presente nella fonte; nessuna voce “involtini” o “bombette” è visibile nella stessa pagina."]),
        ("Contorni", "5Sartoria Gastronomica – Sartoria Gastronomica.pdf", dinner_sides, []),
        ("Insalatone", "6Sartoria Gastronomica – Sartoria Gastronomica.pdf", dinner_salads, []),
        ("Burger", "7Sartoria Gastronomica – Sartoria Gastronomica.pdf", dinner_burgers, ["Tutti i nostri burger hanno pomodoro, insalata e contorno di spicchi di patate fritte.", "Per ogni ingrediente aggiunto è prevista una maggiorazione di 2,50 €"]),
        ("Dessert", "8Sartoria Gastronomica – Sartoria Gastronomica.pdf", dinner_desserts, []),
        ("Bevande", "9Sartoria Gastronomica – Sartoria Gastronomica.pdf", dinner_drinks, []),
        ("Birre", "10Sartoria Gastronomica – Sartoria Gastronomica.pdf", dinner_beers, []),
        ("Carta dei vini", "11Sartoria Gastronomica – Sartoria Gastronomica.pdf", dinner_wines, [
            "La prima scelta l’abbiamo fatta noi, a te ora la seconda mossa.",
            "Quando si sfoglia la carta dei vini si è sempre nella condizione di scegliere.",
            "Ci teniamo che tu sappia che questa carta nasce proprio da alcune scelte che abbiamo fatto per te.",
            "Le cantine che abbiamo selezionato e ti proponiamo, sono per la maggior parte da noi conosciute personalmente e, quando ci è stato possibile, visitate.",
            "Abbiamo dato la precedenza a produttori, soprattutto piccoli, che ricercano la qualità, tutelano il territorio, rispettano la terra e le tradizioni locali, coltivando le proprie vigne con lo stesso amore con cui noi cuciniamo e ci rivolgiamo ai nostri clienti.",
        ]),
    ]),
]


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_labeled(p, label, value):
    if not value:
        return
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BURGUNDY)
    p.add_run(value)


doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(1.6)
sec.bottom_margin = Cm(1.6)
sec.left_margin = Cm(1.65)
sec.right_margin = Cm(1.65)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(9.2)
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = 1.05
for name, size, color, before, after in [
    ("Title", 25, BURGUNDY, 0, 10), ("Heading 1", 17, BURGUNDY, 14, 6),
    ("Heading 2", 13, GOLD, 10, 4), ("Heading 3", 11, BURGUNDY, 6, 2),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

header = sec.header.paragraphs[0]
header.text = "SARTORIA GASTRONOMICA  ·  TRASCRIZIONE INTEGRALE"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in header.runs:
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(GRAY)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("Documento di lavoro — testo originale non tradotto  ·  ")
field = OxmlElement("w:fldSimple")
field.set(qn("w:instr"), "PAGE")
footer._p.append(field)
for r in footer.runs:
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(GRAY)

p = doc.add_paragraph(style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Trascrizione integrale dei menu")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Sartoria Gastronomica · Bologna")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor.from_string(GOLD)

intro = doc.add_paragraph()
intro.add_run("Scopo. ").bold = True
intro.add_run("Documento di lavoro completo per la successiva strutturazione e traduzione in MenuInterattivo. Comprende nomi, descrizioni, ingredienti, prezzi, formati, varianti, allergeni, disponibilità e testi generali visibili nelle 28 fonti PDF. Non costituisce un menu pubblicato e non contiene traduzioni.")
intro = doc.add_paragraph()
intro.add_run("Criterio editoriale. ").bold = True
intro.add_run("Il contenuto è organizzato per menu e categoria; la fonte è indicata in ogni sezione. Refusi, duplicati e incongruenze della fonte sono conservati o segnalati. Logo e barre di navigazione ripetute non sono replicate a ogni pagina; le etichette di categoria sono rappresentate dalla struttura del documento.")

doc.add_heading("Inventario sintetico", level=1)
tbl = doc.add_table(rows=1, cols=3)
tbl.autofit = False
tbl.columns[0].width = Cm(8.2)
tbl.columns[1].width = Cm(3.0)
tbl.columns[2].width = Cm(4.0)
for i, text in enumerate(("Area", "Fonti", "Voci trascritte")):
    c = tbl.rows[0].cells[i]
    c.text = text
    shade(c, BURGUNDY)
    for r in c.paragraphs[0].runs:
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.bold = True
counts = [("Cocktail Bar", "5 PDF", "82"), ("Menù Pranzo", "5 PDF", "37 + 6 extra"), ("Sartoria Bakery Cafè", "7 PDF", "53 (1 duplicato incluso)"), ("Sartoria Gastronomica — cena", "11 PDF", "79"), ("Totale", "28 PDF", "251 schede + 6 extra")]
for row in counts:
    cells = tbl.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
        set_cell_margins(cells[i])
    if row[0] == "Totale":
        for c in cells:
            shade(c, PALE)
            for r in c.paragraphs[0].runs:
                r.bold = True

doc.add_heading("Testi ricorrenti e interfaccia", level=1)
for t in [
    "Marchio ricorrente: SARTORIA GASTRONOMICA.",
    "Titolo ricorrente del menu cena: LA TUA CENA PERFETTA.",
    "Navigazione Cocktail Bar: I CLASSICI · SIGNATURE · GIN & TONIC · BIRRE · STUZZICHERIE.",
    "Navigazione Menù Pranzo: PRIMI · SECONDI · INSALATONE · PANINI GIGANTI · DESSERT.",
    "Navigazione Bakery Cafè: PASTICCERIA · CAFFETTERIA · CAFFETTERIA SPECIALE · BEVANDE ESTIVE · SOFT DRINK · I CLASSICI · BIRRE.",
    "Navigazione cena: ANTIPASTI · PRIMI · TAGLIATE SU MISURA · SECONDI · CONTORNI · INSALATONE · BURGER · DESSERT · BEVANDE · BIRRE · CARTA DEI VINI.",
    "Nota ricorrente su tutte le pagine della cena: Coperto 2,50 €.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(t)

global_index = 0
for area_index, (area, cats) in enumerate(sections):
    doc.add_page_break()
    doc.add_heading(area, level=1)
    for cat, source, products, notes in cats:
        doc.add_heading(cat, level=2)
        p = doc.add_paragraph()
        r = p.add_run(f"Fonte: {source}")
        r.italic = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string(GRAY)
        for note in notes:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.25)
            p.paragraph_format.right_indent = Cm(0.25)
            p.paragraph_format.keep_together = True
            r = p.add_run("Nota della fonte: ")
            r.bold = True
            r.font.color.rgb = RGBColor.from_string(GOLD)
            p.add_run(note)
        for item in products:
            global_index += 1
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.columns[0].width = Cm(13.3)
            table.columns[1].width = Cm(3.2)
            table.rows[0].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table.rows[0].cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for c in table.rows[0].cells:
                set_cell_margins(c, top=45, start=70, bottom=35, end=70)
                shade(c, PALE)
            left, right = table.rows[0].cells
            left.text = ""
            pr = left.paragraphs[0]
            pr.paragraph_format.keep_with_next = True
            r = pr.add_run(f"{global_index}. {item['name']}")
            r.bold = True
            r.font.size = Pt(10.2)
            r.font.color.rgb = RGBColor.from_string(BURGUNDY)
            right.text = item["price"]
            right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in right.paragraphs[0].runs:
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(GOLD)
            if item["description"]:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.2)
                p.paragraph_format.space_after = Pt(1)
                add_labeled(p, "Descrizione / dettagli: ", item["description"])
            if item["allergens"]:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.2)
                p.paragraph_format.space_after = Pt(1)
                add_labeled(p, "Allergeni: ", item["allergens"])
            if item["notes"]:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.2)
                p.paragraph_format.space_after = Pt(3)
                add_labeled(p, "Nota di trascrizione: ", item["notes"])

doc.add_page_break()
doc.add_heading("Punti da approvare prima dell’importazione", level=1)
for text in [
    "Titolo Signature “PEPE AR CU...IWI”: la fonte stessa è troncata/illeggibile; non è stato ricostruito.",
    "Refusi conservati: PENCILLIN, NY CHEESECACKE, CHEESECACKE, CANELLE REPONDE e altre grafie esattamente come mostrate.",
    "Duplicato conservato: MORETTI IPA 5.0% compare due volte nella pagina Birre del Bakery Cafè.",
    "Incongruenza conservata: MORETTI BIONDA è 4.6% nel Cocktail Bar/Bakery e 5.0% nella pagina Birre della cena.",
    "Le righe allergeni sono state trascritte come presenti; eventuali omissioni o classificazioni incongruenti della fonte richiedono revisione del ristorante.",
    "La pagina Tagliate parla di prezzo ad ettogrammo, mentre le due schede mostrano prezzi €/kg: entrambe le formulazioni sono state mantenute.",
    "Nessun contenuto è stato tradotto, normalizzato per la pubblicazione o importato nel progetto.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"Saved {OUT} with {global_index} numbered entries")

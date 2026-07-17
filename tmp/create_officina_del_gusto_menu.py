from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("output/documents/Menu_demo_Officina_del_Gusto_45_prodotti.docx")
INK = "24312D"
GREEN = "315B4A"
COPPER = "B56D3E"
PALE = "F3F0EA"
GRAY = "6B716E"


def p(name, description="", price="", allergens="", note=""):
    return {"name": name, "description": description, "price": price, "allergens": allergens, "note": note}


sections = [
    ("Aperitivi & cocktail", [
        p("Spritz Aperol", "Aperol, prosecco, soda", "6,50 €"),
        p("Hugo Spritz", "Sciroppo di fiore di sambuco, prosecco, menta", "8,00 €"),
        p("Americano", "Vermouth rosso, bitter Campari, soda", "8,00 €"),
        p("Negroni", "Gin, Vermouth rosso, bitter Campari", "9,00 €"),
        p("Paloma", "Tequila, soda al pompelmo, lime", "10,00 €"),
        p("Moscow Mule", "Vodka, lime, ginger beer", "9,00 €"),
        p("Gin Tonic", "Luis Gin 8, tonica Schweppes", "9,00 €"),
        p("Analcolico alla frutta", "Fantasia del barman", "8,00 €"),
    ]),
    ("Antipasti", [
        p("Crudo di Parma", "Stagionato 24 mesi – Selezione Villani", "16,00 €"),
        p("Tagliere di salumi assortiti", "Per 2 persone, con mortadella IGP del consorzio e squacquerone", "14,00 €", "Latte (7)"),
        p("Battuta al coltello di Fassona", "Con uovo fritto e maionese all’acciuga", "25,00 €", "Uova (3), Pesce (4)"),
        p("Carpaccio di manzo", "Con crema di parmigiano e pomodorini confit", "18,00 €", "Latte (7)"),
        p("Burrata con cicoria", "Cicoria ripassata al peperoncino", "16,00 €", "Latte (7)"),
        p("Polpette di melanzane e menta fritte", "Con salsa tzatziki", "14,00 €", "Cereali – Glutine (1), Uova (3), Latte (7)"),
    ]),
    ("Primi", [
        p("Tortellini fatti a mano", "Con ragù bianco su crema di Parmigiano Reggiano", "17,00 €", "Cereali – Glutine (1), Uova (3), Latte (7)"),
        p("Tortellini fatti a mano in brodo", "", "15,00 €", "Cereali – Glutine (1), Uova (3), Latte (7)"),
        p("Tortellino oro di Bologna", "Con crema di panna, parmigiano e zafferano", "16,00 €", "Cereali – Glutine (1), Uova (3), Latte (7)"),
        p("Quadratoni di melanzane, scamorza e ricotta", "Con vellutata di pomodoro e basilico", "16,00 €", "Cereali – Glutine (1), Uova (3), Latte (7)"),
        p("Ravioli ripieni di salsiccia, peperoni e ricotta", "Con ragù di salsiccia fatta da noi", "16,00 €", "Cereali – Glutine (1), Uova (3), Latte (7)"),
    ]),
    ("Secondi & griglia", [
        p("Petto di pollo alla griglia", "Sale grosso e rosmarino con patate al rosmarino", "16,00 €", note="Gluten free"),
        p("Bistecca di manzo alla griglia", "Con patate al rosmarino", "16,00 €", note="Gluten free"),
        p("Filetto di Fassona alla griglia", "", "26,00 €"),
        p("Tagliata New York Top Quality", "Con fonduta di gorgonzola e noci", "25,00 €", "Frutta a guscio (8), Latte (7)"),
        p("Tagliata New York Top Quality", "Con pesto di peperoncini", "24,00 €"),
        p("Costola di suino caramellata", "Con salsa BBQ e miele alla griglia", "18,00 €", "Anidride solforosa e solfiti (12)"),
    ]),
    ("Burger & insalatone", [
        p("Smashed burger di manzo", "Burger di manzo (200 g), cheddar, bacon e salsa burger, con patate fritte", "15,00 €", "Semi di sesamo (11), Cereali – Glutine (1), Latte (7), Uova (3)"),
        p("Bolognese", "Burger di manzo (200 g), crema di parmigiano, mortadella e topping al balsamico", "20,00 €", "Semi di sesamo (11), Cereali – Glutine (1), Latte (7), Uova (3)"),
        p("Piccante", "Burger di manzo (200 g), cheddar, pancetta croccante e salsa piccante", "16,00 €", "Cereali – Glutine (1), Latte (7), Uova (3), Semi di sesamo (11)"),
        p("Insalata Officina pollo", "Insalata mista con petto di pollo alla griglia e scaglie di Parmigiano Reggiano", "15,00 €", "Latte (7)"),
        p("Insalata greca", "Insalata mista con cetrioli, feta, cipolla di Tropea, olive kalamata e origano", "15,00 €", "Latte (7)"),
        p("Panzanella vegana", "Crostini di pane al basilico, pomodorini, olive kalamata e tofu al basilico", "14,00 €", "Cereali – Glutine (1)"),
    ]),
    ("Contorni", [
        p("Patate arrosto al sale grosso e rosmarino", "", "6,00 €"),
        p("Patate fritte", "", "6,00 €"),
        p("Spinaci freschi saltati", "Con olio EVO e pepe", "7,00 €"),
        p("Verdure di stagione alla griglia", "", "8,00 €"),
    ]),
    ("Dessert", [
        p("New York cheesecake", "Ai frutti di bosco o al cioccolato", "8,00 €", "Cereali – Glutine (1), Uova (3), Latte (7)"),
        p("Fondente leggero", "Con crema chantilly e panna montata", "7,00 €", "Latte (7)"),
        p("Tiramisù al bicchiere", "Con mascarpone e caffè", "8,00 €", "Cereali – Glutine (1), Latte (7), Uova (3)"),
        p("Tortino alle nocciole e cioccolato", "Con crema alla gianduia", "7,00 €", "Cereali – Glutine (1), Uova (3), Frutta a guscio (8)"),
        p("Torta di mele e cannella", "Servita con gelato alla crema e panna montata", "8,00 €", "Cereali – Glutine (1), Uova (3), Latte (7)"),
    ]),
    ("Caffetteria & bevande", [
        p("Caffè espresso", "", "2,00 €"),
        p("Cappuccino", "Latte vegetale (soia o avena) +0,50 €", "2,50 €"),
        p("Tè caldo", "", "4,00 €"),
        p("Acqua minerale", "Naturale o frizzante, 0,75 l", "3,00 €"),
        p("Bibite in bottiglia", "33 cl", "4,00 €"),
    ]),
]


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    tc_pr.append(node)


def margins(cell, top=55, start=80, bottom=45, end=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = OxmlElement(f"w:{key}")
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")
        tc_mar.append(element)


def label(paragraph, title, value):
    if not value:
        return
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(GREEN)
    paragraph.add_run(value)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.6)
section.bottom_margin = Cm(1.6)
section.left_margin = Cm(1.7)
section.right_margin = Cm(1.7)

normal = doc.styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(9.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = 1.04
for style_name, size, color in (("Title", 27, GREEN), ("Subtitle", 13, COPPER), ("Heading 1", 18, GREEN), ("Heading 2", 13, COPPER)):
    style = doc.styles[style_name]
    style.font.name = "Aptos Display"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = True
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = header.add_run("OFFICINA DEL GUSTO  ·  MENU DIMOSTRATIVO")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor.from_string(GRAY)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("MenuInterattivo · esempio da approvare · ")
field = OxmlElement("w:fldSimple")
field.set(qn("w:instr"), "PAGE")
footer._p.append(field)
for run in footer.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("OFFICINA DEL GUSTO")
sub = doc.add_paragraph(style="Subtitle")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("Ristorante · Bakery · Cocktail Bar")
intro = doc.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
intro.add_run("Menu dimostrativo da 45 prodotti").bold = True
intro.add_run("\nUna selezione più compatta, costruita a partire dai materiali del menu originale.")

doc.add_paragraph()
box = doc.add_table(rows=1, cols=1)
shade(box.cell(0, 0), PALE)
margins(box.cell(0, 0), 120, 150, 120, 150)
text = box.cell(0, 0).paragraphs[0]
text.alignment = WD_ALIGN_PARAGRAPH.CENTER
text.add_run("Struttura: 8 categorie · 45 prodotti · descrizioni, prezzi e allergeni inclusi\n").bold = True
text.add_run("Coperto 2,50 € · Per allergie o intolleranze rivolgersi al personale.")

doc.add_heading("Indice", level=1)
for name, products in sections:
    row = doc.add_paragraph()
    row.paragraph_format.space_after = Pt(2)
    row.add_run(name).bold = True
    row.add_run(f"  ·  {len(products)} prodotti")

number = 0
for section_index, (name, products) in enumerate(sections):
    doc.add_page_break()
    doc.add_heading(name, level=1)
    for item in products:
        number += 1
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Cm(13.2)
        table.columns[1].width = Cm(3.2)
        left, right = table.rows[0].cells
        for cell in (left, right):
            shade(cell, PALE)
            margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        left.text = ""
        run = left.paragraphs[0].add_run(f"{number}. {item['name']}")
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(GREEN)
        right.text = item["price"]
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in right.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(COPPER)
        if item["description"]:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.2)
            paragraph.paragraph_format.space_after = Pt(1)
            label(paragraph, "Descrizione: ", item["description"])
        if item["allergens"]:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.2)
            paragraph.paragraph_format.space_after = Pt(1)
            label(paragraph, "Allergeni: ", item["allergens"])
        if item["note"]:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.2)
            paragraph.paragraph_format.space_after = Pt(2)
            label(paragraph, "Nota: ", item["note"])

doc.add_page_break()
doc.add_heading("Nota sul documento", level=1)
for text in [
    "Questo è un menu dimostrativo compatto: 45 prodotti contro le 251 schede del materiale originale.",
    "Il nome “Officina del Gusto” è fittizio e sostituisce il nome del ristorante di origine.",
    "I prodotti sono una selezione dei contenuti trascritti; alcune descrizioni sono state rese più uniformi per questo esempio.",
    "Prima di un utilizzo reale, prezzi, ingredienti e allergeni devono essere approvati dal ristorante.",
    "Il documento non è stato importato o pubblicato in MenuInterattivo.",
]:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"Saved {OUT} with {number} products")

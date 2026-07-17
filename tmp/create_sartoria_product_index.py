from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("output/documents/Elenco_completo_prodotti_Sartoria_Gastronomica.docx")

ACCENT = RGBColor(140, 47, 27)  # named override: sartoria_accent
INK = RGBColor(30, 30, 30)
MUTED = RGBColor(96, 96, 96)


menus = [
    (
        "Cocktail Bar",
        [
            (
                "I classici",
                "Cocktail Bar – Sartoria Gastronomica.pdf",
                [
                    "SPRITZ APEROL", "SPRITZ CAMPARI", "HUGO SPRITZ", "AMERICANO",
                    "TOMMY’S MARGARITA", "NEGRONI", "NEGRONI SBAGLIATO", "CAMPARI ORANGE",
                    "MI TO", "CUBA LIBRE", "PALOMA", "GIN TONIC", "GIN TONIC PREMIUM",
                    "MOSCOW MULE", "MOJITO", "BOLOGNA Menegiks Gin & Tonic", "BLOODY MARY",
                    "MANHATTAN", "BOULEVARDIER", "WHISKY SOUR", "DAIQUIRI", "MARGARITA",
                    "MIMOSA", "KIR", "KIR ROYAL", "DRY MARTINI", "GIN FIZZ",
                    "ANALCOLICO FRUTTA", "OLD CUBAN", "OLD FASHIONED", "SAZERAC",
                    "VIEUX CARRE", "WHITE LADY", "WHITE RUSSIAN", "PENCILLIN", "CARDINALE",
                    "HANKY PANKY", "DARK’N’STORMY", "MINT JULEP", "TOM COLLINS", "GOD FATHER",
                    "GOD MOTHER", "TRINIDAD SOUR", "JUNGLE BIRD",
                ],
            ),
            (
                "Signature",
                "2Cocktail Bar – Sartoria Gastronomica.pdf",
                [
                    "CINNAMON CAL", "PEPE AR CU...IWI [titolo illeggibile nella fonte]",
                    "SEDANO DI VASINICOLA", "ARANCIA MOSCATA", "GOLDEN FIZZ",
                    "RED PEPPER LADY", "ANICE ROYALE", "CANELLE REPONDE", "SOFY RAMONES",
                ],
            ),
            (
                "Gin & Tonic",
                "3Cocktail Bar – Sartoria Gastronomica.pdf",
                [
                    "TOSCANA Sabatini & Tonic", "TOSCANA Peter in Florence & Tonic",
                    "TOSCANA Peter in the Navy & Tonic", "TOSCANA 1592 & Tonic",
                    "BASSA MODENA Tabar & Tonic", "BOLOGNA Seawake & Tonic",
                    "BOLOGNA Bad Gin & Tonic", "BOLOGNA Nettuno Blu & Tonic",
                    "BOLOGNA Nettuno Dry & Tonic",
                ],
            ),
            (
                "Birre",
                "4Cocktail Bar – Sartoria Gastronomica.pdf",
                [
                    "MORETTI BIONDA 4.6%", "MORETTI IPA 5.0%", "MORETTI BIANCA 5.0%",
                    "MORETTI 0.0 ALCOOL", "LA GRIGNA PILS", "LATEX PIÙ BLANCHE",
                    "JAKE PACIFIC PALE ALE", "MY BO RED", "AMERICAN HYPE IPA", "WEISS",
                ],
            ),
            (
                "Stuzzicherie",
                "5Cocktail Bar – Sartoria Gastronomica.pdf",
                [
                    "TAGLIERE SALUMI ASSORTITI – PER 2 PERSONE",
                    "TAGLIERE SALUMI ASSORTITI – PER 4 PERSONE", "POLPETTE DI MANZO",
                    "POLPETTE DI PESCE SPADA FRITTE", "POLPETTE MELANZANE E MENTA",
                    "POLPETTE DI LENTICCHIE", "COPPO SICILIANO", "COPPO CALAMARI E GAMBERI",
                    "CESTINO DI SPICCHI DI PIADA", "TIGELLA ARTIGIANALE",
                ],
            ),
        ],
    ),
    (
        "Menù Pranzo",
        [
            (
                "Primi",
                "Menù Pranzo – Sartoria Gastronomica.pdf",
                [
                    "PRIMO DEL GIORNO", "LASAGNA SFOGLIA VERDE ALLA BOLOGNESE",
                    "PARMIGIANA DI MELANZANE", "ZUPPA DI LEGUMI E CEREALI", "PITA GRECA",
                ],
            ),
            (
                "Secondi",
                "2Menù Pranzo – Sartoria Gastronomica.pdf",
                [
                    "PETTO DI POLLO ALLA GRIGLIA", "BISTECCA DI MANZO ALLA GRIGLIA",
                    "PESCE DEL GIORNO", "BOMBETTE DI MAIALE", "SPIEDINI DI POLLO SPEZIATO",
                    "SPIEDINO TURCO", "SPECIAL DEL GIORNO", "MEZZO POLLO AL FORNO",
                ],
            ),
            (
                "Insalatone",
                "3Menù Pranzo – Sartoria Gastronomica.pdf",
                [
                    "INSALATA DI POLLO", "CAESAR SALAD", "INSALATA DI MANZO",
                    "PANZANELLA CROCCANTE", "PANZANELLA CROCCANTE VEG", "SALMONE MARINATO",
                    "GRECA", "PEPERONI RIPIENI", "NOCI E GORGONZOLA", "SPECK, CARCIOFI E OLIVE",
                    "POLLO SFILACCIATO", "TONNO E AVOCADO",
                ],
            ),
            (
                "Extra disponibili",
                "3Menù Pranzo – Sartoria Gastronomica.pdf",
                [
                    "MOZZARELLA DI BUFALA", "FETA", "AVOCADO", "SALMONE MARINATO",
                    "POLLO STRACCETTI", "PANINO SENZA GLUTINE",
                ],
            ),
            (
                "Panini giganti",
                "4Menù Pranzo – Sartoria Gastronomica.pdf",
                [
                    "CLASSICO", "SARTORIA BURGER", "SMASHED BURGER MANZO",
                    "SMASHED BURGER POLLO", "PANINO CON BISTECCA", "VEGETARIANO",
                ],
            ),
            (
                "Dessert",
                "5Menù Pranzo – Sartoria Gastronomica.pdf",
                [
                    "NY CHEESECAKE – CON FRUTTI DI BOSCO", "NY CHEESECAKE – CON CIOCCOLATO",
                    "TENERINA", "TORTA DI MELE", "GELATO ALLA CREMA", "AFFOGATO",
                ],
            ),
        ],
    ),
    (
        "Sartoria Bakery Café",
        [
            (
                "Pasticceria",
                "Sartoria Bakery Cafè – Sartoria Gastronomica.pdf",
                [
                    "PASTICCERIA MIGNON", "CROISSANT", "TENERINA AL CIOCCOLATO", "APPLE PIE",
                    "NY CHEESECACKE [grafia in fonte]", "CHEESECACKE [grafia in fonte]",
                ],
            ),
            (
                "Caffetteria",
                "Sartoria Bakery Cafè – Sartoria Gastronomica_1.pdf",
                ["CAFFÈ ESPRESSO", "CAPPUCCINO", "CAFFÈ AMERICANO", "ORZO", "GINSENG", "TÈ CALDO"],
            ),
            (
                "Caffetteria speciale",
                "2Sartoria Bakery Cafè – Sartoria Gastronomica.pdf",
                [
                    "CREMA CAFFÈ", "NERO DI BRONTE", "CAFFÈ SHAKERATO", "CAFFÈ ESTIVO",
                    "CAFFÈ SALENTINO", "CAFFÈ IN GHIACCIO",
                ],
            ),
            (
                "Bevande estive",
                "3.pdf",
                [
                    "CREMA DI CAFFÈ", "FROZEN YOGURT AL CIOCCOLATO",
                    "FROZEN YOGURT CON AMARENE E MERINGA", "FROZEN YOGURT CON CARAMELLO",
                ],
            ),
            (
                "Soft drink",
                "S5artoria Bakery Cafè – Sartoria Gastronomica.pdf",
                [
                    "SPREMUTA D’ARANCIA", "SUCCHI DI FRUTTA", "BIBITE IN BOTTIGLIA 33 CL",
                    "PICCOLI APERITIVI", "BOLLICINE", "CALICE DI VINO", "SANGRIA",
                ],
            ),
            (
                "I classici",
                "6Sartoria Bakery Cafè – Sartoria Gastronomica.pdf",
                [
                    "SPRITZ APEROL", "SPRITZ CAMPARI", "HUGO SPRITZ", "AMERICANO",
                    "TOMMY’S MARGARITA", "NEGRONI", "NEGRONI SBAGLIATO", "CAMPARI ORANGE",
                    "MI TO", "CUBA LIBRE", "PALOMA", "GIN TONIC", "GIN TONIC PREMIUM",
                    "MOSCOW MULE", "MOJITO",
                ],
            ),
            (
                "Birre",
                "7Sartoria Bakery Cafè – Sartoria Gastronomica.pdf",
                [
                    "MORETTI AL MELOGRANO 5%", "MORETTI BIONDA 4.6%", "MORETTI IPA 5.0%",
                    "MORETTI IPA 5.0% [duplicato nella fonte]", "MORETTI 0.0 ALCOOL",
                    "LA GRIGNA PILS", "LATEX PIÙ BLANCHE", "JAKE PACIFIC PALE ALE", "MY BO RED",
                ],
            ),
        ],
    ),
    (
        "Sartoria Gastronomica – Cena",
        [
            (
                "Antipasti",
                "Sartoria Gastronomica – Sartoria Gastronomica.pdf",
                [
                    "CRUDO DI PARMA", "TAGLIERE DI SALUMI ASSORTITI PER 2 PERSONE",
                    "TAGLIERE DI SALUMI ASSORTITI PER 4 PERSONE",
                    "CESTINO DI SPICCHI DI PIADINE ARTIGIANALI", "TIGELLA ARTIGIANALE",
                    "BATTUTA AL COLTELLO DI FASSONA", "CARPACCIO DI MANZO",
                    "SFORMATINO DI VERDURE", "BURRATA CON CICORIA", "BURRATA CON CRUDO DI PARMA",
                    "SELEZIONE DI SALUMI E FORMAGGI", "SFORMATINO DI PATATE E MORTADELLA",
                    "POLPETTE DI ZUCCA E PATATE", "POLPETTE DI MELANZANE E MENTA FRITTE",
                    "POLPETTE DI LENTICCHIE", "POLPETTE DI PESCE SPADA FRITTE",
                    "COCKTAIL DI GAMBERI",
                ],
            ),
            (
                "Primi",
                "2Sartoria Gastronomica – Sartoria Gastronomica.pdf",
                [
                    "TORTELLINI FATTI A MANO", "TORTELLINI FATTI A MANO IN BRODO",
                    "TORTELLINO ORO DI BOLOGNA", "QUADRATONI DI MELANZANE, SCAMORZA E RICOTTA",
                    "RAVIOLI DI RICOTTA, ’NDUJA E CACIOCAVALLO",
                    "RAVIOLI RIPIENI DI SALSICCIA, PEPERONI E RICOTTA",
                ],
            ),
            (
                "Tagliate su misura",
                "3Sartoria Gastronomica – Sartoria Gastronomica.pdf",
                ["FIORENTINA", "COSTATA"],
            ),
            (
                "Secondi",
                "4Sartoria Gastronomica – Sartoria Gastronomica.pdf",
                [
                    "FILETTO DI FASSONA ALLA GRIGLIA",
                    "TAGLIATA NEW YORK TOP QUALITY – CON FONDUTA DI GORGONZOLA E NOCI",
                    "TAGLIATA NEW YORK TOP QUALITY – CON PESTO DI PEPERONCINI",
                    "TAGLIATA NEW YORK TOP QUALITY – CON SCAGLIE DI PECORINO E GRANELLA DI PISTACCHIO",
                    "TAGLIATA NEW YORK TOP QUALITY – CON LARDO DI PATANEGRA",
                    "COSTOLA DI SUINO CARAMELLATA",
                ],
            ),
            (
                "Contorni",
                "5Sartoria Gastronomica – Sartoria Gastronomica.pdf",
                [
                    "PATATE GRATINATE", "PATATE ARROSTO AL SALE GROSSO E ROSMARINO",
                    "PATATE FRITTE*", "CHIPS DI PATATE SALE E ROSMARINO",
                    "SPINACI FRESCHI SALTATI CON OLIO EVO E PEPE",
                    "VERDURE DI STAGIONE ALLA GRIGLIA", "FUNGHI PORCINI FRITTI*",
                    "CICORIA SALTATA IN PADELLA",
                ],
            ),
            (
                "Insalatone",
                "6Sartoria Gastronomica – Sartoria Gastronomica.pdf",
                [
                    "SPECIAL SARTORIA POLLO", "SPECIAL SARTORIA MANZO", "INSALATA DI SALMONE",
                    "INSALATA AVOCADO E GAMBERI", "INSALATA GRECA", "PANZANELLA VEGANA",
                ],
            ),
            (
                "Burger",
                "7Sartoria Gastronomica – Sartoria Gastronomica.pdf",
                [
                    "SMASHED BURGER DI POLLO", "SMASHED BURGER DI MANZO", "ALTA SARTORIA",
                    "BOLOGNESE", "MODENESE", "PICCANTE",
                ],
            ),
            (
                "Dessert",
                "8Sartoria Gastronomica – Sartoria Gastronomica.pdf",
                [
                    "NEW YORK CHEESECAKE", "FONDENTE LEGGERO", "TIRAMISÙ AL BICCHIERE",
                    "MILLEFOGLIE SCOMPOSTA", "TORTA DI MELE E CANNELLA",
                    "TORTINO ALLE NOCCIOLE E CIOCCOLATO",
                ],
            ),
            (
                "Bevande",
                "9Sartoria Gastronomica – Sartoria Gastronomica.pdf",
                ["ACQUA PANNA", "ACQUA SAN PELLEGRINO", "BIBITE IN BOTTIGLIA"],
            ),
            (
                "Birre",
                "10Sartoria Gastronomica – Sartoria Gastronomica.pdf",
                [
                    "MORETTI BIONDA 5.0%", "MORETTI IPA 5.0%", "MORETTI BIANCA 5.0%",
                    "MORETTI 0.0 ALCOOL", "LA GRIGNA PILS", "LATEX PIÙ BLANCHE",
                    "JAKE PACIFIC PALE ALE", "MY BO RED", "AMERICAN HYPE IPA", "WEISS",
                ],
            ),
            (
                "Carta dei vini",
                "11Sartoria Gastronomica – Sartoria Gastronomica.pdf",
                [
                    "BOLLICINE AL CALICE", "BIANCHI AL CALICE", "ROSSI AL CALICE",
                    "QUINTESSENZA PIGNOLETTO FRIZZANTE MEDICI ERMETE",
                    "PROSECCO SUPERIORE DI VALDOBBIADENE DOCG BRUT BELLENDA",
                    "FRANCIACORTA BRUT DOCG FERGHETTINA",
                    "FRANCIACORTA ROSÉ DOCG FERGHETTINA",
                    "PIGNOLETTO PERDITEMPO TERRE ROSSE", "LUGANA CÀ DEI FRATI",
                ],
            ),
        ],
    ),
]


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Pagina ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_numbering_definition(doc):
    numbering = doc.part.numbering_part.element
    existing = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "right")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "left")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)
    return abstract_id


def new_number_id(doc, abstract_id):
    numbering = doc.part.numbering_part.element
    existing = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def number_paragraph(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_ref = OxmlElement("w:numId")
    num_ref.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_ref)
    p_pr.append(num_pr)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, before, after in (
        ("Heading 1", 16, 18, 10),
        ("Heading 2", 13, 14, 7),
        ("Heading 3", 12, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    run = header_p.add_run("Sartoria Gastronomica · Elenco completo dei prodotti")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_cover(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(42)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    run = kicker.add_run("MENUINTERATTIVO · DOCUMENTO DI LAVORO")
    set_run_font(run, size=10, color=ACCENT, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Elenco completo dei prodotti")
    set_run_font(run, size=30, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    run = subtitle.add_run("Sartoria Gastronomica · Bologna")
    set_run_font(run, size=16, color=ACCENT)

    description = doc.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    description.paragraph_format.space_after = Pt(28)
    run = description.add_run(
        "Tutti i nomi estratti dai 28 PDF, organizzati per menu e categoria. "
        "Il documento contiene soltanto i nomi: nessuna traduzione, descrizione, prezzo o allergene."
    )
    set_run_font(run, size=11, color=MUTED)

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.paragraph_format.space_after = Pt(5)
    run = summary.add_run("251 schede prodotto · 6 extra pranzo · 28 categorie")
    set_run_font(run, size=12, color=INK, bold=True)

    summary2 = doc.add_paragraph()
    summary2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary2.paragraph_format.space_after = Pt(22)
    run = summary2.add_run("Totale righe nel documento: 257")
    set_run_font(run, size=10.5, color=MUTED)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(20)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run(
        "Le diciture tra parentesi quadre identificano anomalie già presenti nelle fonti."
    )
    set_run_font(run, size=9.5, color=MUTED, italic=True)
    doc.add_page_break()


def add_overview(doc):
    doc.add_heading("Indice sintetico", level=1)
    overview = [
        ("Cocktail Bar", 82, 5),
        ("Menù Pranzo", 43, 6),
        ("Sartoria Bakery Café", 53, 7),
        ("Sartoria Gastronomica – Cena", 79, 11),
    ]
    for label, products, categories in overview:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r1 = p.add_run(f"{label}: ")
        set_run_font(r1, bold=True, color=INK)
        r2 = p.add_run(f"{products} righe, {categories} categorie")
        set_run_font(r2, color=INK)

    doc.add_heading("Criterio editoriale", level=1)
    for text in (
        "I nomi sono mantenuti nella forma visibile nei PDF, senza correzioni silenziose.",
        "Le categorie seguono l’ordine della navigazione originale.",
        "La numerazione è progressiva e globale, da 1 a 257, per facilitare i riferimenti.",
        "Il duplicato Moretti IPA del Bakery Café è riportato due volte e segnalato.",
        "Il titolo Signature non leggibile integralmente resta marcato come tale.",
    ):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("•  " + text)
        set_run_font(r, color=INK)
    doc.add_page_break()


def build_document():
    product_rows = sum(len(items) for _, categories in menus for _, _, items in categories)
    assert product_rows == 257, product_rows
    assert sum(len(items) for _, categories in menus for title, _, items in categories if title != "Extra disponibili") == 251
    assert [sum(len(items) for _, _, items in categories) for _, categories in menus] == [82, 43, 53, 79]

    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
    abstract_id = add_numbering_definition(doc)

    core = doc.core_properties
    core.title = "Elenco completo dei prodotti – Sartoria Gastronomica"
    core.subject = "Estratto dei nomi prodotto dai 28 PDF forniti"
    core.author = "MenuInterattivo"
    core.keywords = "Sartoria Gastronomica, menu, prodotti, categorie"

    add_cover(doc)
    add_overview(doc)

    for menu_index, (menu_name, categories) in enumerate(menus):
        h1 = doc.add_heading(menu_name, level=1)
        h1.paragraph_format.keep_with_next = True

        total = sum(len(items) for _, _, items in categories)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(f"{total} righe · {len(categories)} categorie")
        set_run_font(r, size=9.5, color=MUTED, italic=True)

        for category_name, source, items in categories:
            h2 = doc.add_heading(f"{category_name} ({len(items)})", level=2)
            h2.paragraph_format.keep_with_next = True
            source_p = doc.add_paragraph()
            source_p.paragraph_format.space_before = Pt(0)
            source_p.paragraph_format.space_after = Pt(5)
            source_p.paragraph_format.keep_with_next = True
            source_run = source_p.add_run(f"Fonte: {source}")
            set_run_font(source_run, size=8.5, color=MUTED, italic=True)

            num_id = new_number_id(doc, abstract_id)
            for item in items:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.25
                number_paragraph(p, num_id)
                r = p.add_run(item)
                set_run_font(r, size=10.5, color=INK)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build_document()
